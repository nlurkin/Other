#!/bin/env python
# encoding: utf-8
'''
FillClaim -- shortdesc

FillClaim is a description

It defines classes_and_methods

@author:     user_name

@copyright:  2017 organization_name. All rights reserved.

@license:    license

@contact:    user_email
@deffield    updated: Updated
'''

import sys

from argparse import ArgumentParser
from argparse import RawDescriptionHelpFormatter
from openpyxl import Workbook, load_workbook
from docx import Document

__all__ = []
__version__ = 0.1
__date__ = '2017-06-01'
__updated__ = '2017-06-01'

class xlsImportInfo():
    def __init__(self):
        self.totalVAT = []
        self.totalWVAT = []
        self.total = []
        
        self.travel = []
        self.travelVAT = []
        self.travelByCar = []
        self.travelByCarVAT = []
        self.accomodation = []
        self.accomodationVAT = []
        self.subsistence = []
        self.subsistenceVAT = []
        self.taxi = []
        self.taxiVAT = []
        self.other = []
        self.otherVAT = []
    
    def __repr__(self):
        return str(self.__dict__)

def readXlsx(filePath):
    values = xlsImportInfo()
    wb = load_workbook(filePath, data_only=True)
    totalSheet = wb['Total']
    values.totalVAT         = [str(int(round(totalSheet['F19'].value))),str(int(round(totalSheet['G19'].value)))]
    values.totalWVAT        = [str(int(round(totalSheet['C19'].value))),str(int(round(totalSheet['D19'].value)))]
    values.total            = [str(int(round(totalSheet['C20'].value))),str(int(round(totalSheet['D20'].value)))]
    values.travel           = [str(int(round(totalSheet['C12'].value))),str(int(round(totalSheet['D12'].value)))]
    values.travelVAT        = [str(int(round(totalSheet['F12'].value))),str(int(round(totalSheet['G12'].value)))]
    values.travelByCar      = [str(int(round(totalSheet['C13'].value))),str(int(round(totalSheet['D13'].value)))]
    values.travelByCarVAT   = [str(int(round(totalSheet['F13'].value))),str(int(round(totalSheet['G13'].value)))]
    values.accomodation     = [str(int(round(totalSheet['C14'].value))),str(int(round(totalSheet['D14'].value)))]
    values.accomodationVAT  = [str(int(round(totalSheet['F14'].value))),str(int(round(totalSheet['G14'].value)))]
    values.subsistence      = [str(int(round(totalSheet['C15'].value))),str(int(round(totalSheet['D15'].value)))]
    values.subsistenceVAT   = [str(int(round(totalSheet['F15'].value))),str(int(round(totalSheet['G15'].value)))]
    values.taxi             = [str(int(round(totalSheet['C16'].value))),str(int(round(totalSheet['D16'].value)))]
    values.taxiVAT          = [str(int(round(totalSheet['F16'].value))),str(int(round(totalSheet['G16'].value)))]
    values.other            = [str(int(round(totalSheet['C17'].value))),str(int(round(totalSheet['D17'].value)))]
    values.otherVAT         = [str(int(round(totalSheet['F17'].value))),str(int(round(totalSheet['G17'].value)))]
    
    print values
    return values

def fillVal(doc, amountIndex, penceIndex, values):
    amountCol = doc.tables[0].column_cells(amountIndex[0])
    penceCol = doc.tables[0].column_cells(penceIndex[0])
    print values
    if values[0]!="0" or values[1]!="0":
        [amountCol[amountIndex[1]].text ,penceCol[penceIndex[1]].text ] = values
    else:
        [amountCol[amountIndex[1]].text ,penceCol[penceIndex[1]].text ] = ["",""]

def findField(doc, name, col_range, row_range):
    amount_index = (None, None)
    pence_index = (None, None)
    for col in col_range:
        amountCol = doc.tables[0].column_cells(col)
        for row in row_range:
            if row>=len(amountCol):
                break
            val = amountCol[row].text
            if val=="{"+name+"}":
                amount_index = (col, row)
            if val=="{"+name+"_p}":
                pence_index = (col, row)

            #if amount_index != (None,None) and pence_index != (None, None):
            #    break

        #if amount_index != (None,None) and pence_index != (None, None):
        #    break
        
    if amount_index == (None,None) or pence_index == (None, None):
        print "Error: " + name + " not found"
        sys.exit(1)

    print "Found " + name + " at ", amount_index, pence_index
    return amount_index, pence_index
    
def modifyDocx(filePath, values):
    doc = Document(filePath)
    
    a,p = findField(doc, "total_tf", range(30, 40), range(5, 10))
    fillVal(doc, a, p, values.totalWVAT)
    a,p = findField(doc, "total_vat", range(30, 40), range(5, 20))
    fillVal(doc, a, p, values.totalVAT)
    a,p = findField(doc, "total", range(30, 40), range(10, 20))
    fillVal(doc, a, p, values.total)

    a,p = findField(doc, "travel", range(30, 40), range(10, 20))
    fillVal(doc, a, p, values.travel)
    a,p = findField(doc, "car", range(30, 40), range(10, 20))
    fillVal(doc, a, p, values.travelByCar)
    a,p = findField(doc, "accom", range(30, 40), range(10, 20))
    fillVal(doc, a, p, values.accomodation)
    a,p = findField(doc, "accom_vat", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.accomodationVAT)
    a,p = findField(doc, "food", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.subsistence)
    a,p = findField(doc, "food_vat", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.subsistenceVAT)
    a,p = findField(doc, "taxi", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.taxi)
    a,p = findField(doc, "taxi_vat", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.taxiVAT)
    a,p = findField(doc, "fuel", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.other)
    a,p = findField(doc, "fuel_vat", range(30, 40), range(20, 30))
    fillVal(doc, a, p, values.otherVAT)
    a,p = findField(doc, "total_b", range(30, 40), range(30, 40))
    fillVal(doc, a, p, values.total)
    
    doc.save("test2.docx")
    
def main():
    '''Command line options.'''

    # Setup argument parser
    parser = ArgumentParser(description="", formatter_class=RawDescriptionHelpFormatter)
    parser.add_argument("-d", "--doc",  help="doc file", required=True)
    parser.add_argument("-x", "--xls", help="excel file", required=True)
    parser.add_argument('-V', '--version', action='version', version=__version__)

    # Process arguments
    args = parser.parse_args()

    values = readXlsx(args.xls)
    modifyDocx(args.doc, values)
    return 0

if __name__ == "__main__":
    sys.exit(main())
